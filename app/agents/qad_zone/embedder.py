"""QAD-Zone document embedder.

After a Word doc is generated for a custom module, the user can choose
to embed it in Qdrant (collection: qad_custom_docs) so Apex can answer
questions about it.

Chunking strategy:
  - Parse the generated .docx file
  - Split into semantic chunks (one per heading section)
  - Each chunk tagged with: module_name, program_name (if extractable),
    chunk_type (section heading), doc_title, source="qad_custom_docs"
  - Upsert to Qdrant collection ``qad_custom_docs``
"""
from __future__ import annotations

import asyncio
import logging
import re
import uuid
from functools import partial
from pathlib import Path
from typing import Any

from docx import Document

from app.core.config import settings
from app.core.llm import openai_embed
from app.vector.qdrant import get_qdrant

logger = logging.getLogger(__name__)

COLLECTION = "qad_custom_docs"
DOWNLOADS_DIR = Path("app/static/downloads")

# Minimum characters for a chunk to be worth embedding
_MIN_CHUNK = 80
# Maximum characters per chunk (roughly 300 tokens)
_MAX_CHUNK = 1200


def _extract_module_name(title: str) -> str:
    """Best-effort module name from doc title, e.g. 'RTDC Delivery Challan Management' → 'rtdc'."""
    if not title:
        return "unknown"
    # Take first word, lowercase
    first = title.strip().split()[0].lower()
    # Strip common suffixes
    first = re.sub(r"[^a-z0-9]", "", first)
    return first or "unknown"


def _read_docx_chunks(doc_path: Path, doc_title: str, module_name: str) -> list[dict[str, Any]]:
    """Parse a .docx file and return list of {text, metadata} chunks."""
    try:
        doc = Document(str(doc_path))
    except Exception as exc:
        logger.error("Cannot open docx %s: %s", doc_path, exc)
        return []

    chunks: list[dict[str, Any]] = []
    current_heading = "Overview"
    current_text_parts: list[str] = []

    def _flush(heading: str, parts: list[str]) -> None:
        text = "\n".join(parts).strip()
        if len(text) < _MIN_CHUNK:
            return
        # Split oversized chunks
        while len(text) > _MAX_CHUNK:
            chunks.append({
                "text": text[:_MAX_CHUNK],
                "metadata": {
                    "module": module_name,
                    "section": heading,
                    "doc_title": doc_title,
                    "source": COLLECTION,
                },
            })
            text = text[_MAX_CHUNK:]
        if len(text) >= _MIN_CHUNK:
            chunks.append({
                "text": text,
                "metadata": {
                    "module": module_name,
                    "section": heading,
                    "doc_title": doc_title,
                    "source": COLLECTION,
                },
            })

    for para in doc.paragraphs:
        style = (para.style.name or "").lower()
        text = para.text.strip()
        if not text:
            continue

        if style.startswith("heading"):
            _flush(current_heading, current_text_parts)
            current_heading = text
            current_text_parts = []
        else:
            current_text_parts.append(text)

    _flush(current_heading, current_text_parts)

    # Also pull text from tables
    for table in doc.tables:
        table_lines: list[str] = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
        if table_lines:
            table_text = "\n".join(table_lines)
            if len(table_text) >= _MIN_CHUNK:
                chunks.append({
                    "text": table_text[:_MAX_CHUNK],
                    "metadata": {
                        "module": module_name,
                        "section": "Table Data",
                        "doc_title": doc_title,
                        "source": COLLECTION,
                    },
                })

    return chunks


def _ensure_collection(client, vector_size: int = 3072) -> None:
    """Create qad_custom_docs collection if it doesn't exist."""
    from qdrant_client.models import Distance, VectorParams
    existing = [c.name for c in client.get_collections().collections]
    if COLLECTION not in existing:
        logger.info("Creating Qdrant collection '%s'", COLLECTION)
        client.create_collection(
            collection_name=COLLECTION,
            vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
        )
    # Ensure payload index on 'module' for fast filtering
    try:
        from qdrant_client.models import PayloadSchemaType
        client.create_payload_index(
            collection_name=COLLECTION,
            field_name="module",
            field_schema=PayloadSchemaType.KEYWORD,
        )
    except Exception:
        pass  # Index may already exist


async def embed_document(doc_url: str, doc_title: str) -> dict[str, Any]:
    """
    Given a doc_url like '/static/downloads/abcd1234.docx' and a title,
    extract text chunks, embed them, and upsert to Qdrant qad_custom_docs.

    Returns {"chunks_embedded": N, "module": module_name}
    """
    # Resolve path
    filename = Path(doc_url).name
    doc_path = DOWNLOADS_DIR / filename
    if not doc_path.exists():
        raise FileNotFoundError(f"Document not found: {doc_path}")

    module_name = _extract_module_name(doc_title)
    logger.info("Embedding doc '%s' as module='%s' (%s)", doc_title, module_name, filename)

    chunks = _read_docx_chunks(doc_path, doc_title, module_name)
    if not chunks:
        raise ValueError("No embeddable content found in document.")

    logger.info("Extracted %d chunks from '%s'", len(chunks), doc_title)

    # Embed all chunks (sequential to avoid rate-limiting)
    vectors: list[list[float]] = []
    for ch in chunks:
        vec = await openai_embed(ch["text"])
        vectors.append(vec)

    # Upsert to Qdrant
    client = get_qdrant()
    loop = asyncio.get_event_loop()

    # Ensure collection exists with correct dimensions
    vec_size = len(vectors[0]) if vectors else 3072
    await loop.run_in_executor(None, partial(_ensure_collection, client, vec_size))

    from qdrant_client.models import PointStruct
    points = [
        PointStruct(
            id=str(uuid.uuid4()),
            vector=vec,
            payload={**ch["metadata"], "text": ch["text"]},
        )
        for ch, vec in zip(chunks, vectors)
    ]

    await loop.run_in_executor(
        None,
        partial(client.upsert, collection_name=COLLECTION, points=points),
    )

    logger.info("Upserted %d points to Qdrant collection '%s'", len(points), COLLECTION)
    return {"chunks_embedded": len(points), "module": module_name}
