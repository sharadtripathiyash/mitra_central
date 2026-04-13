"""Corporate Word document generator for QAD Custom Module Documentation.

Template structure (matches generate_doc.js / template_schema.json):
  Title Page → TOC placeholder → Executive Summary → Architecture →
  Database Tables → Program Analysis → Workflow → Setup Instructions →
  QAD Native Comparison → Error Messages → Customization History →
  Quick Reference → End Page

Design rules:
- Only render a section / subsection / table / bullet if LLM returned real data.
- isEmpty() / _has() guard every piece of content — nothing is ever assumed present.
- Flowchart page is rendered only when FLOWCHART.SHOW=true and cairosvg is available.
"""
from __future__ import annotations

import uuid
import tempfile
import os
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

DOWNLOADS_DIR = Path("app/static/downloads")

# ── Color palette (mirrors generate_doc.js C constants) ──────────────────────
C = {
    "BLUE":     RGBColor(0x1F, 0x4E, 0x79),
    "MBLUE":    RGBColor(0x2E, 0x75, 0xB6),
    "ACCENT":   RGBColor(0xC0, 0x00, 0x00),
    "MGREEN":   RGBColor(0x70, 0xAD, 0x47),
    "DKGREEN":  RGBColor(0x37, 0x56, 0x23),
    "DGRAY":    RGBColor(0x40, 0x40, 0x40),
    "MGRAY":    RGBColor(0x76, 0x76, 0x76),
    "WHITE":    RGBColor(0xFF, 0xFF, 0xFF),
    "YELLOW":   RGBColor(0x85, 0x64, 0x04),
    "TEAL":     RGBColor(0x0C, 0x54, 0x60),
}
FILL = {
    "BLUE":     "1F4E79",
    "MBLUE":    "2E75B6",
    "LBLUE":    "DEEAF1",
    "LGREEN":   "E2EFDA",
    "LGRAY":    "F2F2F2",
    "WHITE":    "FFFFFF",
    "YELLOW":   "FFF3CD",
    "TEAL":     "D1ECF1",
    "CODERED":  "FCE4E4",
    "ACCENT":   "C00000",
}
FONT = "Arial"
ALT  = [FILL["WHITE"], FILL["LGRAY"]]


# ── Helpers ───────────────────────────────────────────────────────────────────

def _has(value: Any) -> bool:
    """Return True only when value carries real, non-empty content."""
    if value is None:
        return False
    if isinstance(value, bool):
        return value
    if isinstance(value, (list, dict)):
        return len(value) > 0
    if isinstance(value, str):
        v = value.strip()
        return bool(v) and v not in ("", "AUTO")
    return bool(value)


def _shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _borders(cell, color: str = "AAAAAA", size: str = "4") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        tcB.append(el)
    tcPr.append(tcB)


def _cell_margins(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcM = OxmlElement("w:tcMar")
    for side, val in [("top", 80), ("bottom", 80), ("left", 120), ("right", 120)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcM.append(el)
    tcPr.append(tcM)


def _set_cell(cell, text: str, fill: str = FILL["WHITE"], bold: bool = False,
              italic: bool = False, color: RGBColor | None = None,
              hdr_color: str | None = None, font_size: int = 10) -> None:
    _shading(cell, fill)
    _borders(cell, color=hdr_color or "AAAAAA")
    _cell_margins(cell)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(str(text) if text is not None else "")
    run.font.name = FONT
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or C["DGRAY"]


def _hdr_cell(cell, text: str, fill: str = FILL["MBLUE"]) -> None:
    _shading(cell, fill)
    _borders(cell, color=fill)
    _cell_margins(cell)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = C["WHITE"]


def _add_header_footer(doc: Document, system_name: str, system_full: str) -> None:
    for section in doc.sections:
        # Header
        hdr = section.header
        hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear()
        hp.paragraph_format.space_after = Pt(3)
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
        bot.set(qn("w:color"), FILL["MBLUE"]); bot.set(qn("w:space"), "1")
        pBdr.append(bot); pPr.append(pBdr)
        r1 = hp.add_run(f"{system_name} — {system_full}")
        r1.font.name = FONT; r1.font.size = Pt(9); r1.font.color.rgb = C["MBLUE"]
        r2 = hp.add_run("    QAD ERP | Progress 4GL Custom Development")
        r2.font.name = FONT; r2.font.size = Pt(9); r2.font.color.rgb = C["MGRAY"]

        # Footer
        ftr = section.footer
        ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear()
        fp.paragraph_format.space_before = Pt(3)
        pPr2 = fp._p.get_or_add_pPr()
        pBdr2 = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "4")
        top.set(qn("w:color"), FILL["MBLUE"]); top.set(qn("w:space"), "1")
        pBdr2.append(top); pPr2.append(pBdr2)
        r3 = fp.add_run("CONFIDENTIAL — Internal Use Only    |    Page ")
        r3.font.name = FONT; r3.font.size = Pt(8); r3.font.color.rgb = C["MGRAY"]
        fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
        page_r = OxmlElement("w:r")
        page_r.append(fldChar1); page_r.append(instr); page_r.append(fldChar2)
        fp._p.append(page_r)


# ── Element builders ──────────────────────────────────────────────────────────

def _h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(18); run.font.bold = True
        run.font.color.rgb = C["BLUE"]
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(14); run.font.bold = True
        run.font.color.rgb = C["MBLUE"]
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "3"); bot.set(qn("w:color"), FILL["MBLUE"])
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)


def _h3(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = FONT; run.font.size = Pt(12); run.font.bold = True
        run.font.color.rgb = C["DGRAY"]
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)


def _para(doc: Document, text: str, italic: bool = False,
          color: RGBColor | None = None, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(size)
    run.font.italic = italic; run.font.color.rgb = color or C["DGRAY"]


def _bullet(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(11); run.font.bold = bold
    run.font.color.rgb = C["DGRAY"]


def _num_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(11); run.font.color.rgb = C["DGRAY"]


def _mono(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "360"); pPr.append(ind)
    run = p.add_run(text)
    run.font.name = "Courier New"; run.font.size = Pt(9); run.font.color.rgb = C["ACCENT"]


def _note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), FILL["YELLOW"])
    pPr.append(shd)
    r1 = p.add_run("⚠ NOTE: "); r1.font.name = FONT; r1.font.size = Pt(10)
    r1.font.bold = True; r1.font.color.rgb = C["YELLOW"]
    r2 = p.add_run(text); r2.font.name = FONT; r2.font.size = Pt(10)
    r2.font.color.rgb = C["YELLOW"]


def _info(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), FILL["TEAL"])
    pPr.append(shd)
    r1 = p.add_run("ℹ INFO: "); r1.font.name = FONT; r1.font.size = Pt(10)
    r1.font.bold = True; r1.font.color.rgb = C["TEAL"]
    r2 = p.add_run(text); r2.font.name = FONT; r2.font.size = Pt(10)
    r2.font.color.rgb = C["TEAL"]


def _divider(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4"); bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot); pPr.append(pBdr)


def _data_table(doc: Document, headers: list[str], rows: list[list]) -> None:
    """Render a styled table. Skips automatically if rows is empty."""
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = "Table Grid"
    hrow = table.add_row()
    for i, h in enumerate(headers):
        _hdr_cell(hrow.cells[i], h)
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        for ci in range(len(headers)):
            val = row_data[ci] if ci < len(row_data) else ""
            _set_cell(row.cells[ci], str(val) if val is not None else "", fill=ALT[ri % 2])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _kv_table(doc: Document, pairs: list[tuple[str, str]]) -> None:
    """Two-column label/value table."""
    if not pairs:
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in pairs:
        row = table.add_row()
        _hdr_cell(row.cells[0], label, fill=FILL["MBLUE"])
        _set_cell(row.cells[1], value, fill=FILL["LBLUE"], italic=True, color=C["MBLUE"])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Section builders ──────────────────────────────────────────────────────────

def _build_title_page(doc: Document, T: dict, doc_date: str) -> None:
    """Cover / title page."""
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(48)

    if _has(T.get("SYSTEM_NAME")):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(T["SYSTEM_NAME"])
        r.font.name = FONT; r.font.size = Pt(48); r.font.bold = True
        r.font.color.rgb = C["BLUE"]

    if _has(T.get("SYSTEM_FULL_NAME")):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(T["SYSTEM_FULL_NAME"])
        r.font.name = FONT; r.font.size = Pt(22); r.font.color.rgb = C["MBLUE"]

    # Accent divider
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_before = Pt(8); div.paragraph_format.space_after = Pt(8)
    pPr = div._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:color"), FILL["ACCENT"])
    pBdr.append(bot); pPr.append(pBdr)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Comprehensive Technical Documentation")
    r2.font.name = FONT; r2.font.size = Pt(14); r2.font.color.rgb = C["DGRAY"]

    if _has(T.get("PLATFORM")):
        plat = doc.add_paragraph()
        plat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = plat.add_run(T["PLATFORM"])
        r3.font.name = FONT; r3.font.size = Pt(12); r3.font.color.rgb = C["MGRAY"]

    doc.add_paragraph().paragraph_format.space_before = Pt(20)

    # Metadata table
    meta_pairs: list[tuple[str, str]] = []
    for label, key in [
        ("System Name",      "SYSTEM_FULL_NAME"),
        ("Platform",         "PLATFORM"),
        ("Module",           "MODULE"),
        ("Version",          "VERSION"),
        ("Original Author",  "ORIGINAL_AUTHOR"),
        ("Last Modified By", "LAST_MODIFIED_BY"),
        ("Total Programs",   "TOTAL_PROGRAMS"),
    ]:
        val = T.get(key)
        if _has(val):
            meta_pairs.append((label, val))
    meta_pairs.append(("Document Date", doc_date))
    _kv_table(doc, meta_pairs)

    doc.add_page_break()


def _build_toc_placeholder(doc: Document) -> None:
    """TOC placeholder (update manually in Word via right-click → Update Field)."""
    _h1(doc, "Table of Contents")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), FILL["LGRAY"])
    pPr.append(shd)
    run = p.add_run("[ Right-click here and select 'Update Field' to populate the Table of Contents ]")
    run.font.name = FONT; run.font.size = Pt(10); run.font.italic = True
    run.font.color.rgb = C["MGRAY"]
    doc.add_page_break()


def _build_section1(doc: Document, ES: dict) -> None:
    """1. Executive Summary"""
    _h1(doc, "1.  Executive Summary")

    if _has(ES.get("PARA_1")):
        _para(doc, ES["PARA_1"])
    if _has(ES.get("PARA_2")):
        _para(doc, ES["PARA_2"])

    if _has(ES.get("KEY_CAPABILITIES")):
        _h2(doc, "1.1  Key Business Capabilities")
        for cap in ES["KEY_CAPABILITIES"]:
            if _has(cap):
                _bullet(doc, cap)

    ct = ES.get("COMPARISON_TABLE") or {}
    if _has(ct.get("rows")) and _has(ct.get("headers")):
        _h2(doc, "1.2  Relationship to Standard QAD Functionality")
        _data_table(doc, ct["headers"], ct["rows"])

    doc.add_page_break()


def _build_section2(doc: Document, ARCH: dict) -> None:
    """2. System Architecture & Program Inventory"""
    _h1(doc, "2.  System Architecture & Program Inventory")

    if _has(ARCH.get("INTRO_PARA")):
        _para(doc, ARCH["INTRO_PARA"])

    ph = ARCH.get("PROGRAM_HIERARCHY_TABLE") or {}
    if _has(ph.get("rows")) and _has(ph.get("headers")):
        _h2(doc, "2.1  Program Hierarchy")
        _data_table(doc, ph["headers"], ph["rows"])

    sv = ARCH.get("SHARED_VARIABLES_TABLE") or {}
    if _has(sv.get("rows")) and _has(sv.get("headers")):
        _h2(doc, "2.2  Shared Variable Architecture")
        _data_table(doc, sv["headers"], sv["rows"])

    doc.add_page_break()


def _build_section3(doc: Document, tables: list) -> None:
    """3. Database Tables"""
    _h1(doc, "3.  Database Tables — Structure & Unique Keys")

    for t in tables:
        if not isinstance(t, dict):
            continue
        name = t.get("TABLE_NAME", "")
        subtitle = t.get("TABLE_SUBTITLE", "")
        heading = f"{name}" + (f" — {subtitle}" if _has(subtitle) else "")
        _h3(doc, heading)

        if _has(t.get("TABLE_DESCRIPTION")):
            _para(doc, t["TABLE_DESCRIPTION"])

        tf = t.get("TABLE_FIELDS") or {}
        if _has(tf.get("rows")) and _has(tf.get("headers")):
            _data_table(doc, tf["headers"], tf["rows"])

        if _has(t.get("TABLE_UNIQUE_KEY")):
            p = doc.add_paragraph()
            r = p.add_run(f"Unique Key: {t['TABLE_UNIQUE_KEY']}")
            r.font.name = FONT; r.font.size = Pt(11); r.font.bold = True
            r.font.color.rgb = C["DGRAY"]

        if _has(t.get("TABLE_NOTE")):
            _note(doc, t["TABLE_NOTE"])
        if _has(t.get("TABLE_INFO")):
            _info(doc, t["TABLE_INFO"])

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.add_page_break()


def _build_section4(doc: Document, programs: list) -> None:
    """4. Detailed Program Analysis"""
    _h1(doc, "4.  Detailed Program Analysis")

    for idx, prog in enumerate(programs, 1):
        if not isinstance(prog, dict):
            continue

        prog_name = prog.get("PROG_NAME", f"Program {idx}")
        _h2(doc, f"4.{idx}  {prog_name}")

        if _has(prog.get("PROG_VERSION_INFO")):
            _para(doc, prog["PROG_VERSION_INFO"], italic=True, color=C["MGRAY"], size=9)

        # Purpose
        if _has(prog.get("PROG_PURPOSE")):
            _h3(doc, "Purpose")
            _para(doc, prog["PROG_PURPOSE"])

        # Call graph
        meta: list[tuple[str, str]] = []
        if _has(prog.get("PROG_CALLED_BY")):
            meta.append(("Called By", prog["PROG_CALLED_BY"]))
        if _has(prog.get("PROG_CALLS")):
            calls = prog["PROG_CALLS"]
            meta.append(("Calls", ", ".join(calls) if isinstance(calls, list) else str(calls)))
        if _has(prog.get("PROG_INCLUDE_FILES")):
            inc = prog["PROG_INCLUDE_FILES"]
            meta.append(("Include Files", ", ".join(inc) if isinstance(inc, list) else str(inc)))
        if meta:
            _kv_table(doc, meta)

        # Screen layout
        sl = prog.get("PROG_SCREEN_LAYOUT") or {}
        if _has(sl.get("rows")) and _has(sl.get("headers")):
            frame = sl.get("FRAME_NAME", "")
            _h3(doc, f"Screen Layout{(' (' + frame + ')') if _has(frame) else ''}")
            _data_table(doc, sl["headers"], sl["rows"])

        # Logic steps
        if _has(prog.get("PROG_LOGIC_STEPS")):
            _h3(doc, "Key Processing Logic")
            for step in prog["PROG_LOGIC_STEPS"]:
                if _has(step):
                    _num_item(doc, step)

        # Validations
        if _has(prog.get("PROG_VALIDATIONS")):
            _h3(doc, "Key Validations")
            for v in prog["PROG_VALIDATIONS"]:
                if _has(v):
                    _bullet(doc, v)

        # Triggers
        if _has(prog.get("PROG_TRIGGERS")):
            _h3(doc, "WRITE / ASSIGN Triggers")
            for t in prog["PROG_TRIGGERS"]:
                if _has(t):
                    _bullet(doc, t)

        # Special tables
        sp = prog.get("PROG_SPECIAL_TABLES") or {}
        if sp.get("SHOW") and _has(sp.get("rows")) and _has(sp.get("headers")):
            _h3(doc, "Availability / Special Reference")
            _data_table(doc, sp["headers"], sp["rows"])

        # Extra section
        ex = prog.get("PROG_EXTRA_SECTION") or {}
        if ex.get("SHOW") and _has(ex.get("TITLE")):
            _h3(doc, ex["TITLE"])
            ct = ex.get("CONTENT_TYPE", "")
            if ct == "para" and _has(ex.get("PARA")):
                _para(doc, ex["PARA"])
            elif ct == "bullets" and _has(ex.get("BULLETS")):
                for b in ex["BULLETS"]:
                    if _has(b):
                        _bullet(doc, b)
            elif ct == "table":
                tbl = ex.get("TABLE") or {}
                if _has(tbl.get("rows")) and _has(tbl.get("headers")):
                    _data_table(doc, tbl["headers"], tbl["rows"])

    doc.add_page_break()


def _build_section5(doc: Document, WF: dict) -> None:
    """5. Business Workflow & Process Flow"""
    _h1(doc, "5.  Business Workflow & Process Flow")

    if _has(WF.get("INTRO_PARA")):
        _para(doc, WF["INTRO_PARA"])

    ph = WF.get("PHASES_TABLE") or {}
    if _has(ph.get("rows")) and _has(ph.get("headers")):
        _h2(doc, "5.1  High-Level Process Flow")
        _data_table(doc, ph["headers"], ph["rows"])

    if _has(WF.get("INTERNAL_CALL_FLOW")):
        _h2(doc, "5.2  Internal Program Call Flow")
        for line in WF["INTERNAL_CALL_FLOW"]:
            _mono(doc, line)

    aw = WF.get("APPROVAL_WORKFLOW") or {}
    if aw.get("SHOW") and _has(aw.get("STEPS")):
        _h2(doc, "5.3  Approval Workflow Detail")
        for step in aw["STEPS"]:
            if _has(step):
                _num_item(doc, step)
        if _has(aw.get("NOTE")):
            _note(doc, aw["NOTE"])

    dr = WF.get("DELETE_RULES_TABLE") or {}
    if _has(dr.get("rows")) and _has(dr.get("headers")):
        _h2(doc, "5.4  Delete Rules")
        _data_table(doc, dr["headers"], dr["rows"])

    doc.add_page_break()


def _build_section6(doc: Document, SETUP: dict) -> None:
    """6. Setup & Deployment Instructions"""
    _h1(doc, "6.  Setup & Deployment Instructions")

    if _has(SETUP.get("PREREQUISITES")):
        _h2(doc, "6.1  Prerequisites")
        for p in SETUP["PREREQUISITES"]:
            if _has(p):
                _bullet(doc, p)

    steps = SETUP.get("STEPS") or []
    for i, step in enumerate(steps):
        if not isinstance(step, dict):
            continue
        num = step.get("STEP_NUMBER", str(i + 1))
        title = step.get("STEP_TITLE", "")
        _h2(doc, f"6.{i + 2}  Step {num}" + (f" — {title}" if _has(title) else ""))
        if _has(step.get("STEP_DESCRIPTION")):
            _para(doc, step["STEP_DESCRIPTION"])
        if _has(step.get("STEP_ITEMS")):
            for item in step["STEP_ITEMS"]:
                if _has(item):
                    _num_item(doc, item)
        if _has(step.get("STEP_CODE")):
            for line in step["STEP_CODE"]:
                _mono(doc, line)

    mt = SETUP.get("MENU_TABLE") or {}
    if mt.get("SHOW") and _has(mt.get("rows")) and _has(mt.get("headers")):
        _h2(doc, f"6.{len(steps) + 2}  Menu Configuration")
        _data_table(doc, mt["headers"], mt["rows"])

    if _has(SETUP.get("TEST_STEPS")):
        _h2(doc, f"6.{len(steps) + 3}  End-to-End Test")
        for s in SETUP["TEST_STEPS"]:
            if _has(s):
                _num_item(doc, s)

    doc.add_page_break()


def _build_section7(doc: Document, NAT: dict) -> None:
    """7. QAD Native Functionality (optional, SHOW flag)"""
    if not (NAT and NAT.get("SHOW")):
        return
    _h1(doc, "7.  QAD Native Functionality — Comparison & Setup")

    if _has(NAT.get("NATIVE_DESCRIPTION_PARA")):
        _para(doc, NAT["NATIVE_DESCRIPTION_PARA"])

    mods = NAT.get("NATIVE_MODULES") or []
    for i, m in enumerate(mods):
        if not isinstance(m, dict):
            continue
        if _has(m.get("MODULE_NAME")):
            _h3(doc, f"7.1.{i + 1}  {m['MODULE_NAME']}")
        if _has(m.get("MODULE_DESCRIPTION")):
            _para(doc, m["MODULE_DESCRIPTION"])

    native_steps = NAT.get("NATIVE_SETUP_STEPS") or []
    if native_steps:
        _h2(doc, "7.2  How to Set Up in Native QAD")
        for s in native_steps:
            if not isinstance(s, dict):
                continue
            if _has(s.get("STEP_TITLE")):
                _h3(doc, s["STEP_TITLE"])
            for item in (s.get("STEP_ITEMS") or []):
                if _has(item):
                    _num_item(doc, item)

    dd = NAT.get("DEPLOYMENT_DECISION_TABLE") or {}
    if _has(dd.get("rows")) and _has(dd.get("headers")):
        _h2(doc, "7.3  Deployment Decision Matrix")
        _data_table(doc, dd["headers"], dd["rows"])

    doc.add_page_break()


def _build_section8(doc: Document, ERR: dict) -> None:
    """8. Error Messages & Handling"""
    tbl = ERR.get("TABLE") or {}
    if not (_has(tbl.get("rows")) and _has(tbl.get("headers"))):
        return
    _h1(doc, "8.  Error Messages & Handling")
    _data_table(doc, tbl["headers"], tbl["rows"])
    doc.add_page_break()


def _build_section9(doc: Document, history: list) -> None:
    """9. Customization History & Enhancement Notes"""
    if not history:
        return
    _h1(doc, "9.  Customization History & Enhancement Notes")

    for i, eco in enumerate(history, 1):
        if not isinstance(eco, dict):
            continue
        eco_id = eco.get("ECO_ID", "")
        eco_title = eco.get("ECO_TITLE", "")
        heading = eco_id + (f" — {eco_title}" if _has(eco_title) else "")
        _h2(doc, f"9.{i}  {heading}")

        author = eco.get("ECO_AUTHOR", "")
        date = eco.get("ECO_DATE", "")
        meta_line = "  |  ".join(filter(_has, [
            f"Author: {author}" if _has(author) else "",
            f"Date: {date}" if _has(date) else "",
        ]))
        if meta_line:
            _para(doc, meta_line, italic=True, color=C["MGRAY"], size=9)

        for change in (eco.get("ECO_CHANGES") or []):
            if _has(change):
                _bullet(doc, change)

    doc.add_page_break()


def _build_section10(doc: Document, QR: dict) -> None:
    """10. Quick Reference Card"""
    # Determine if there is anything to render at all
    ref_tables = [
        (QR.get("TRANSACTION_TYPE_TABLE"), "10.1  Transaction Type Guide"),
        (QR.get("AUTH_GROUP_TABLE"),       "10.2  Authorization Group Flow"),
        (QR.get("INCLUDE_FILES_TABLE"),    "10.3  Key Include Files"),
        (QR.get("LOT_SERIAL_TABLE"),       "10.4  Lot / Serial Control Reference"),
        (QR.get("CUSTOM_TABLE_1"),         "10.5  Additional Reference"),
    ]
    any_visible = any(
        t and t.get("SHOW") and _has(t.get("rows")) and _has(t.get("headers"))
        for t, _ in ref_tables
    )
    if not any_visible:
        return

    _h1(doc, "10.  Quick Reference Card")

    for tbl, default_title in ref_tables:
        if not (tbl and tbl.get("SHOW") and _has(tbl.get("rows")) and _has(tbl.get("headers"))):
            continue
        title = tbl.get("TITLE") or default_title
        _h2(doc, title)
        _data_table(doc, tbl["headers"], tbl["rows"])


def _build_end_page(doc: Document, system_name: str, full_name: str, doc_date: str) -> None:
    _divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Document —")
    r.font.name = FONT; r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = C["MGRAY"]

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{system_name} — {full_name}  |  QAD ERP Progress 4GL  |  {doc_date}")
    r2.font.name = FONT; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


# ── Flowchart (optional — requires cairosvg) ─────────────────────────────────

def _hex(h: str):
    """Convert hex string to RGB tuple."""
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _try_build_flowchart(doc: Document, FC: dict, system_name: str) -> None:
    """Render flowchart as PNG using Pillow (pure Python, no Cairo needed).
    Silently skipped if SHOW is false or nodes are missing/invalid.
    """
    if not (FC and FC.get("SHOW")):
        return

    # Filter out any instruction strings the LLM may have left in arrays
    lanes  = [l for l in (FC.get("LANES",  []) or []) if isinstance(l, dict) and l.get("LANE_ID")]
    nodes  = [n for n in (FC.get("NODES",  []) or []) if isinstance(n, dict) and n.get("ID")]
    arrows = [a for a in (FC.get("ARROWS", []) or []) if isinstance(a, dict) and a.get("FROM")]

    if not nodes:
        return

    try:
        import io as _io
        from PIL import Image, ImageDraw, ImageFont

        # ── Canvas setup ──────────────────────────────────────────────────────
        W, H   = 1800, max(900, 120 + len(nodes) * 95)
        n_lanes = max(len(lanes), 1)
        LANE_W  = W // n_lanes

        COLORS = {
            "dark_blue":   {"fill": (0x1F,0x4E,0x79), "text": (255,255,255), "stroke": (0x1F,0x4E,0x79)},
            "light_blue":  {"fill": (0xDE,0xEA,0xF1), "text": (0x1F,0x4E,0x79), "stroke": (0x2E,0x75,0xB6)},
            "green":       {"fill": (0x37,0x56,0x23), "text": (255,255,255), "stroke": (0x37,0x56,0x23)},
            "light_green": {"fill": (0xE2,0xEF,0xDA), "text": (0x37,0x56,0x23), "stroke": (0x70,0xAD,0x47)},
            "red":         {"fill": (0xFC,0xE4,0xD6), "text": (0xC0,0x00,0x00), "stroke": (0xC0,0x00,0x00)},
            "yellow":      {"fill": (0xFF,0xF2,0xCC), "text": (0x7D,0x4E,0x00), "stroke": (0xF4,0xA8,0x00)},
            "gray":        {"fill": (0xF2,0xF2,0xF2), "text": (0x40,0x40,0x40), "stroke": (0x7F,0x7F,0x7F)},
        }
        LANE_BG = {
            "dark_blue":  (0xF7,0xFA,0xFD),
            "light_blue": (0xEB,0xF3,0xFB),
            "green":      (0xF0,0xF7,0xEE),
        }
        LANE_HDR = {
            "dark_blue":  (0x1F,0x4E,0x79),
            "light_blue": (0x2E,0x75,0xB6),
            "green":      (0x37,0x56,0x23),
        }
        ARROW_COLORS = {
            "blue":  (0x1F,0x4E,0x79),
            "green": (0x37,0x56,0x23),
            "red":   (0xC0,0x00,0x00),
        }

        img = Image.new("RGB", (W, H), (0xF7,0xFA,0xFD))
        draw = ImageDraw.Draw(img)

        # Load fonts (fall back to default if not found)
        try:
            fnt_bold  = ImageFont.truetype("arialbd.ttf",  18)
            fnt_norm  = ImageFont.truetype("arial.ttf",    14)
            fnt_small = ImageFont.truetype("arial.ttf",    12)
            fnt_title = ImageFont.truetype("arialbd.ttf",  22)
        except Exception:
            fnt_bold = fnt_norm = fnt_small = fnt_title = ImageFont.load_default()

        # ── Title bar ─────────────────────────────────────────────────────────
        draw.rectangle([0, 0, W, 50], fill=(0x1F,0x4E,0x79))
        title_text = f"{system_name} — System Process Flow"
        draw.text((W // 2, 25), title_text, font=fnt_title, fill=(255,255,255), anchor="mm")

        # ── Lane backgrounds & headers ────────────────────────────────────────
        lane_x: dict[str, int] = {}
        for i, lane in enumerate(lanes):
            lx = i * LANE_W
            lane_x[lane["LANE_ID"]] = lx
            bg  = LANE_BG.get(lane.get("LANE_COLOR",""), (0xF7,0xFA,0xFD))
            hdr = LANE_HDR.get(lane.get("LANE_COLOR",""), (0x1F,0x4E,0x79))
            draw.rectangle([lx, 50, lx + LANE_W, H], fill=bg)
            if i > 0:
                draw.line([lx, 50, lx, H], fill=(0xC5,0xD9,0xEF), width=1)
            draw.rectangle([lx, 50, lx + LANE_W, 82], fill=hdr)
            label = lane.get("LANE_LABEL","").replace("\\n", "\n")
            draw.text((lx + LANE_W // 2, 66), label, font=fnt_small,
                      fill=(255,255,255), anchor="mm", align="center")

        # ── Compute node positions ────────────────────────────────────────────
        lane_count: dict[str, int] = {}
        node_pos: dict[str, tuple] = {}
        for n in nodes:
            lane_id = n.get("LANE", lanes[0]["LANE_ID"] if lanes else "")
            lx = lane_x.get(lane_id, 0)
            idx = lane_count.get(lane_id, 0)
            lane_count[lane_id] = idx + 1
            cx = lx + LANE_W // 2
            cy = 110 + idx * 95
            node_pos[n["ID"]] = (cx, cy)

        NW, NH = LANE_W - 30, 40

        # ── Draw arrows first (behind nodes) ─────────────────────────────────
        for a in arrows:
            frm = node_pos.get(a.get("FROM",""))
            to  = node_pos.get(a.get("TO",""))
            if not frm:
                continue
            tx = to[0] if to else frm[0] + 40
            ty = to[1] if to else frm[1] + 30
            color = ARROW_COLORS.get(a.get("COLOR","blue"), ARROW_COLORS["blue"])
            draw.line([frm[0], frm[1], tx, ty], fill=color, width=2)
            # Arrowhead
            import math
            angle = math.atan2(ty - frm[1], tx - frm[0])
            aw = 10
            ax1 = tx - aw * math.cos(angle - 0.4)
            ay1 = ty - aw * math.sin(angle - 0.4)
            ax2 = tx - aw * math.cos(angle + 0.4)
            ay2 = ty - aw * math.sin(angle + 0.4)
            draw.polygon([(tx,ty),(ax1,ay1),(ax2,ay2)], fill=color)
            if _has(a.get("LABEL")):
                mx = (frm[0] + tx) // 2 + 4
                my = (frm[1] + ty) // 2 - 10
                draw.text((mx, my), a["LABEL"], font=fnt_small, fill=color)

        # ── Draw nodes ────────────────────────────────────────────────────────
        for n in nodes:
            pos = node_pos.get(n.get("ID",""))
            if not pos:
                continue
            cx, cy = pos
            col   = COLORS.get(n.get("COLOR",""), COLORS["dark_blue"])
            label = n.get("LABEL","").replace("{{SYSTEM_NAME}}", system_name).replace("\\n", "\n")
            ntype = n.get("TYPE","box")

            if ntype == "oval":
                rx, ry = NW // 2 - 5, 20
                draw.ellipse([cx-rx, cy-ry, cx+rx, cy+ry], fill=col["fill"], outline=col["stroke"], width=2)
                draw.text((cx, cy), label, font=fnt_bold, fill=col["text"], anchor="mm", align="center")

            elif ntype == "diamond":
                dw, dh = NW // 2, 26
                pts = [(cx, cy-dh), (cx+dw, cy), (cx, cy+dh), (cx-dw, cy)]
                draw.polygon(pts, fill=col["fill"], outline=col["stroke"])
                draw.text((cx, cy), label, font=fnt_small, fill=col["text"], anchor="mm", align="center")

            else:  # box
                x0, y0 = cx - NW//2, cy - NH//2
                draw.rounded_rectangle([x0, y0, x0+NW, y0+NH], radius=5,
                                       fill=col["fill"], outline=col["stroke"], width=2)
                draw.text((cx, cy), label, font=fnt_norm, fill=col["text"], anchor="mm", align="center")

        # ── Save to bytes ─────────────────────────────────────────────────────
        buf = _io.BytesIO()
        img.save(buf, format="PNG")
        png_data = buf.getvalue()

        # ── Embed in Word doc ─────────────────────────────────────────────────
        doc.add_page_break()
        _h1(doc, "System Process Flow")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(_io.BytesIO(png_data), width=Inches(9))
        doc.add_page_break()

    except Exception as e:
        import logging as _log
        _log.getLogger(__name__).warning("Flowchart render skipped: %s", e)


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_document(title: str, sections: list[dict], *, subtitle: str = "Mitra Central") -> str:
    DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)

    # Extract structured data (new template schema)
    D: dict[str, Any] = {}
    for s in sections:
        if s.get("heading") == "structured_data" and s.get("metadata"):
            D = s["metadata"]
            break

    # Legacy fallback for old flat-key responses
    if not D:
        D = {"TITLE_PAGE": {"SYSTEM_NAME": title, "SYSTEM_FULL_NAME": title}}
        for s in sections:
            content = (s.get("content") or "").strip()
            if not content:
                continue
            h = (s.get("heading") or "").lower()
            es = D.setdefault("EXECUTIVE_SUMMARY", {})
            if "purpose" in h:
                es["PARA_1"] = content
            elif "background" in h or "context" in h:
                es["PARA_2"] = content

    # Resolve title-page data
    TP = D.get("TITLE_PAGE") or {}
    sys_name = TP.get("SYSTEM_NAME") or title
    sys_full = TP.get("SYSTEM_FULL_NAME") or title
    doc_date = (
        datetime.now().strftime("%d %B %Y")
        if not _has(TP.get("DOCUMENT_DATE")) or TP.get("DOCUMENT_DATE") == "AUTO"
        else TP["DOCUMENT_DATE"]
    )

    doc = Document()

    # Page setup (A4: 595pt × 842pt in twips = 11906 × 16838 dxa)
    for section in doc.sections:
        section.page_width    = Pt(595)
        section.page_height   = Pt(842)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(0.9)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    _add_header_footer(doc, sys_name, sys_full)

    # ── Build sections ────────────────────────────────────────────────────────
    _build_title_page(doc, TP, doc_date)
    _build_toc_placeholder(doc)

    ES = D.get("EXECUTIVE_SUMMARY")
    if _has(ES):
        _build_section1(doc, ES)

    ARCH = D.get("ARCHITECTURE")
    if _has(ARCH):
        _build_section2(doc, ARCH)

    DB_TABLES = D.get("DATABASE_TABLES")
    if _has(DB_TABLES):
        _build_section3(doc, DB_TABLES)

    PROG = D.get("PROGRAM_ANALYSIS")
    if _has(PROG):
        _build_section4(doc, PROG)

    WF = D.get("WORKFLOW")
    if _has(WF):
        _build_section5(doc, WF)

    SETUP = D.get("SETUP_INSTRUCTIONS")
    if _has(SETUP):
        _build_section6(doc, SETUP)

    NAT = D.get("QAD_NATIVE_COMPARISON")
    if _has(NAT):
        _build_section7(doc, NAT)

    ERR = D.get("ERROR_MESSAGES")
    if _has(ERR):
        _build_section8(doc, ERR)

    HIST = D.get("CUSTOMIZATION_HISTORY")
    if _has(HIST):
        _build_section9(doc, HIST)

    QR = D.get("QUICK_REFERENCE")
    if _has(QR):
        _build_section10(doc, QR)

    FC = D.get("FLOWCHART")
    if _has(FC):
        _try_build_flowchart(doc, FC, sys_name)

    _build_end_page(doc, sys_name, sys_full, doc_date)

    filename = f"{uuid.uuid4().hex[:12]}.docx"
    doc.save(str(DOWNLOADS_DIR / filename))
    return f"/static/downloads/{filename}"
