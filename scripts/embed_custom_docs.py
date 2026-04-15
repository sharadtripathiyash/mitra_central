"""
embed_custom_docs.py
====================
One-time script to (re)build the qad_custom_docs Qdrant collection.

What it does:
  1. Deletes existing 'qad_custom_docs' collection (if present)
  2. Creates a fresh collection with same name
  3. Embeds all source code from 3 ZIPs:
       RTDC.zip, MRN.zip, DOA.zip
       from C:\\Users\\sharad.tripathi\\Downloads\\Customizations
       Files embedded: .p  .i  .df  .xml  (Power Automate JSON inside nested zips too)
  4. Embeds the 3 Word documentation files from app/static/downloads/
       MRN_System_Documentation.docx
       DOA_System_Documentation.docx
       RTDC_System_Documentation.docx

Credentials: read from .env in project root (QDRANT_URL, QDRANT_API_KEY, OPENAI_API_KEY)
Collection:  hardcoded as "qad_custom_docs" (matches apex/service.py CUSTOM_DOCS_COLLECTION)

Run from project root:
    python scripts/embed_custom_docs.py
"""
from __future__ import annotations

import json
import os
import re
import sys
import tempfile
import time
import uuid
import zipfile
from pathlib import Path
from typing import Any

# ── Load .env from project root ───────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent.parent / ".env")
except ImportError:
    print("[warn] python-dotenv not installed — reading env vars directly")

import httpx
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, FieldCondition, Filter, MatchAny, PayloadSchemaType, PointStruct, VectorParams

# ── Config ────────────────────────────────────────────────────────────────────
COLLECTION      = "qad_custom_docs"          # hardcoded — matches apex/service.py
VECTOR_SIZE     = 3072                        # text-embedding-3-large

QDRANT_URL      = os.environ["QDRANT_URL"]
QDRANT_API_KEY  = os.environ.get("QDRANT_API_KEY", "")
OPENAI_API_KEY  = os.environ["OPENAI_API_KEY"]
EMBED_MODEL     = os.environ.get("OPENAI_EMBED_MODEL", "text-embedding-3-large")

ZIPS_DIR        = Path(r"C:\Users\sharad.tripathi\Downloads\Customizations")
DOCS_DIR        = Path(__file__).parent.parent / "app" / "static" / "downloads"

# Map zip filename → module key
ZIP_MODULE_MAP = {
    "MRN.zip":  "mrn",
    "DOA.zip":  "doa",
    "RTDC.zip": "rtdc",
}

# Map docx filename → module key
DOCX_MODULE_MAP = {
    "MRN_System_Documentation.docx":  "mrn",
    "DOA_System_Documentation.docx":  "doa",
    "RTDC_System_Documentation.docx": "rtdc",
}

# Files to embed from ZIPs
CODE_EXTENSIONS = {".p", ".i", ".df", ".xml"}

# Chunk size limits
MIN_CHUNK = 80
MAX_CHUNK = 1200

# ── OpenAI embed ──────────────────────────────────────────────────────────────
def openai_embed(text: str) -> list[float]:
    """Call OpenAI embeddings API synchronously."""
    resp = httpx.post(
        "https://api.openai.com/v1/embeddings",
        headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
        json={"input": text[:8000], "model": EMBED_MODEL},
        timeout=30,
    )
    resp.raise_for_status()
    return resp.json()["data"][0]["embedding"]


# ── Qdrant setup ──────────────────────────────────────────────────────────────
def setup_collection(client: QdrantClient) -> None:
    existing = [c.name for c in client.get_collections().collections]
    if COLLECTION in existing:
        print(f"  Deleting existing collection '{COLLECTION}' ...")
        client.delete_collection(COLLECTION)
        time.sleep(1)

    print(f"  Creating collection '{COLLECTION}' (size={VECTOR_SIZE}, cosine) ...")
    client.create_collection(
        collection_name=COLLECTION,
        vectors_config=VectorParams(size=VECTOR_SIZE, distance=Distance.COSINE),
    )

    # Payload index on 'module' for fast filtering
    try:
        client.create_payload_index(
            collection_name=COLLECTION,
            field_name="module",
            field_schema=PayloadSchemaType.KEYWORD,
        )
        print("  Created payload index on 'module'")
    except Exception as exc:
        print(f"  [warn] Could not create payload index: {exc}")


# ── Chunking helpers ──────────────────────────────────────────────────────────

def _chunk_by_size(text: str, max_size: int = MAX_CHUNK) -> list[str]:
    """Split text into chunks of max_size, respecting newlines where possible."""
    chunks = []
    while len(text) > max_size:
        # Try to split at a newline near max_size
        split_at = text.rfind("\n", 0, max_size)
        if split_at < max_size // 2:
            split_at = max_size
        chunks.append(text[:split_at].strip())
        text = text[split_at:].strip()
    if len(text) >= MIN_CHUNK:
        chunks.append(text)
    return chunks


def chunk_progress_code(content: str, filename: str, module: str) -> list[dict[str, Any]]:
    """
    Chunk a Progress 4GL .p or .i file by PROCEDURE / FUNCTION blocks.
    Falls back to size-based chunking if no blocks found.
    """
    chunks: list[dict[str, Any]] = []

    # Add a short file-level overview chunk
    header_lines = []
    for line in content.splitlines()[:30]:
        stripped = line.strip()
        if stripped.startswith("/*") or stripped.startswith("*") or stripped.startswith("//"):
            header_lines.append(stripped)
        elif stripped:
            header_lines.append(stripped)
        if len(header_lines) >= 15:
            break

    if header_lines:
        header_text = f"File: {filename} (Module: {module.upper()})\n" + "\n".join(header_lines)
        chunks.append({
            "text": header_text,
            "metadata": {
                "module": module, "filename": filename,
                "file_type": "program", "chunk_type": "file_overview",
                "source": COLLECTION,
            }
        })

    # Split on PROCEDURE / FUNCTION / CLASS boundaries
    pattern = re.compile(
        r'(?:^|\n)((?:PROCEDURE|FUNCTION|CLASS|METHOD)\s+\S.*?)(?=\n(?:PROCEDURE|FUNCTION|CLASS|METHOD)\s|\Z)',
        re.IGNORECASE | re.DOTALL
    )
    matches = list(pattern.finditer(content))

    if matches:
        for m in matches:
            block = m.group(1).strip()
            if len(block) < MIN_CHUNK:
                continue
            # Determine chunk_type
            first_word = block.split()[0].upper() if block else "BLOCK"
            # Extract name
            name_match = re.match(r'(?:PROCEDURE|FUNCTION|CLASS|METHOD)\s+(\S+)', block, re.IGNORECASE)
            block_name = name_match.group(1).strip(":").strip('"') if name_match else "unknown"

            for sub in _chunk_by_size(block):
                chunks.append({
                    "text": f"[{filename} / {first_word}: {block_name}]\n{sub}",
                    "metadata": {
                        "module": module, "filename": filename,
                        "file_type": "program", "chunk_type": first_word.lower(),
                        "block_name": block_name, "source": COLLECTION,
                    }
                })
    else:
        # No procedure blocks — chunk by size
        for sub in _chunk_by_size(content):
            chunks.append({
                "text": f"[{filename}]\n{sub}",
                "metadata": {
                    "module": module, "filename": filename,
                    "file_type": "program", "chunk_type": "code_block",
                    "source": COLLECTION,
                }
            })

    return chunks


def chunk_df_file(content: str, filename: str, module: str) -> list[dict[str, Any]]:
    """
    Chunk a Progress .df (data dictionary) file by ADD TABLE blocks.
    Each table + its fields becomes one chunk.
    """
    chunks: list[dict[str, Any]] = []
    table_pattern = re.compile(r'(ADD TABLE\s+"?\w+"?.*?)(?=ADD TABLE|\Z)', re.IGNORECASE | re.DOTALL)

    matches = list(table_pattern.finditer(content))
    if matches:
        for m in matches:
            block = m.group(1).strip()
            if len(block) < MIN_CHUNK:
                continue
            # Extract table name
            tname_match = re.search(r'ADD TABLE\s+"?(\w+)"?', block, re.IGNORECASE)
            table_name = tname_match.group(1) if tname_match else "unknown"

            for sub in _chunk_by_size(block):
                chunks.append({
                    "text": f"[{filename} / Table: {table_name} (Module: {module.upper()})]\n{sub}",
                    "metadata": {
                        "module": module, "filename": filename, "table_name": table_name,
                        "file_type": "schema", "chunk_type": "table_definition",
                        "source": COLLECTION,
                    }
                })
    else:
        for sub in _chunk_by_size(content):
            chunks.append({
                "text": f"[{filename}]\n{sub}",
                "metadata": {
                    "module": module, "filename": filename,
                    "file_type": "schema", "chunk_type": "schema_block",
                    "source": COLLECTION,
                }
            })
    return chunks


def chunk_xml_file(content: str, filename: str, module: str) -> list[dict[str, Any]]:
    """Chunk XML files — usually small controller/metadata files."""
    chunks: list[dict[str, Any]] = []
    for sub in _chunk_by_size(content):
        chunks.append({
            "text": f"[{filename} / XML Metadata (Module: {module.upper()})]\n{sub}",
            "metadata": {
                "module": module, "filename": filename,
                "file_type": "xml", "chunk_type": "xml_metadata",
                "source": COLLECTION,
            }
        })
    return chunks


def chunk_json_flow(content: str, filename: str, module: str) -> list[dict[str, Any]]:
    """Chunk Power Automate / cloud flow JSON files by action entries."""
    chunks: list[dict[str, Any]] = []
    try:
        data = json.loads(content)
        # Try to extract flow properties
        flow_name = data.get("name") or data.get("properties", {}).get("displayName", filename)
        description = data.get("properties", {}).get("description", "")
        # Extract actions if present
        actions = {}
        try:
            actions = data["properties"]["definition"]["actions"]
        except (KeyError, TypeError):
            pass

        if description:
            chunks.append({
                "text": f"[Power Automate Flow: {flow_name} (Module: {module.upper()})]\n{description}",
                "metadata": {
                    "module": module, "filename": filename,
                    "file_type": "cloud_flow", "chunk_type": "flow_overview",
                    "source": COLLECTION,
                }
            })

        if actions:
            for action_name, action_body in actions.items():
                action_text = json.dumps({action_name: action_body}, indent=2)
                if len(action_text) >= MIN_CHUNK:
                    for sub in _chunk_by_size(action_text):
                        chunks.append({
                            "text": f"[Flow: {flow_name} / Action: {action_name}]\n{sub}",
                            "metadata": {
                                "module": module, "filename": filename,
                                "file_type": "cloud_flow", "chunk_type": "flow_action",
                                "action_name": action_name, "source": COLLECTION,
                            }
                        })
        else:
            # Fallback: dump full JSON in chunks
            for sub in _chunk_by_size(json.dumps(data, indent=2)):
                chunks.append({
                    "text": f"[{filename}]\n{sub}",
                    "metadata": {
                        "module": module, "filename": filename,
                        "file_type": "cloud_flow", "chunk_type": "flow_definition",
                        "source": COLLECTION,
                    }
                })
    except json.JSONDecodeError:
        # Not valid JSON — treat as plain text
        for sub in _chunk_by_size(content):
            chunks.append({
                "text": f"[{filename}]\n{sub}",
                "metadata": {
                    "module": module, "filename": filename,
                    "file_type": "cloud_flow", "chunk_type": "flow_text",
                    "source": COLLECTION,
                }
            })
    return chunks


# ── Extract chunks from a single file ────────────────────────────────────────
def extract_file_chunks(file_path: Path, filename: str, module: str) -> list[dict[str, Any]]:
    """Read a file and return chunks based on its type."""
    ext = file_path.suffix.lower()
    try:
        content = file_path.read_text(encoding="utf-8", errors="replace").strip()
    except Exception as exc:
        print(f"    [warn] Cannot read {filename}: {exc}")
        return []

    if not content or len(content) < MIN_CHUNK:
        return []

    if ext in (".p", ".i"):
        return chunk_progress_code(content, filename, module)
    elif ext == ".df":
        return chunk_df_file(content, filename, module)
    elif ext == ".xml":
        return chunk_xml_file(content, filename, module)
    elif ext == ".json":
        return chunk_json_flow(content, filename, module)
    else:
        # Generic text chunking
        chunks = []
        for sub in _chunk_by_size(content):
            chunks.append({
                "text": f"[{filename}]\n{sub}",
                "metadata": {
                    "module": module, "filename": filename,
                    "file_type": ext.lstrip(".") or "text", "chunk_type": "text_block",
                    "source": COLLECTION,
                }
            })
        return chunks


# ── Process ZIP files ─────────────────────────────────────────────────────────
def collect_chunks_from_zips() -> list[dict[str, Any]]:
    all_chunks: list[dict[str, Any]] = []

    for zip_name, module in ZIP_MODULE_MAP.items():
        zip_path = ZIPS_DIR / zip_name
        if not zip_path.exists():
            print(f"  [warn] ZIP not found: {zip_path} — skipping")
            continue

        print(f"\n  Processing {zip_name} → module='{module}'")
        with tempfile.TemporaryDirectory() as tmpdir:
            with zipfile.ZipFile(zip_path, "r") as zf:
                zf.extractall(tmpdir)

            tmp = Path(tmpdir)
            file_count = 0

            # Walk all extracted files
            for file_path in sorted(tmp.rglob("*")):
                if not file_path.is_file():
                    continue

                ext = file_path.suffix.lower()

                # Handle nested zip (e.g. Power Automate export inside zip)
                if ext == ".zip":
                    print(f"    Found nested ZIP: {file_path.name} — extracting...")
                    nested_dir = file_path.parent / (file_path.stem + "_extracted")
                    nested_dir.mkdir(exist_ok=True)
                    try:
                        with zipfile.ZipFile(file_path, "r") as nzf:
                            nzf.extractall(nested_dir)
                        # Process extracted nested files
                        for nf in sorted(nested_dir.rglob("*")):
                            if nf.is_file() and nf.suffix.lower() in CODE_EXTENSIONS | {".json"}:
                                chunks = extract_file_chunks(nf, nf.name, module)
                                if chunks:
                                    print(f"      {nf.name} → {len(chunks)} chunks")
                                    all_chunks.extend(chunks)
                                    file_count += 1
                    except Exception as exc:
                        print(f"    [warn] Could not process nested zip: {exc}")
                    continue

                # Only process known code file types
                if ext not in CODE_EXTENSIONS | {".json"}:
                    continue

                chunks = extract_file_chunks(file_path, file_path.name, module)
                if chunks:
                    print(f"    {file_path.name} → {len(chunks)} chunks")
                    all_chunks.extend(chunks)
                    file_count += 1

            print(f"  Done {zip_name}: {file_count} files → {sum(1 for c in all_chunks if c['metadata']['module'] == module)} chunks so far")

    return all_chunks


# ── Process Word documentation files ─────────────────────────────────────────
def collect_chunks_from_docx() -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError:
        print("  [warn] python-docx not installed — skipping Word doc embedding")
        return []

    all_chunks: list[dict[str, Any]] = []

    for docx_name, module in DOCX_MODULE_MAP.items():
        doc_path = DOCS_DIR / docx_name
        if not doc_path.exists():
            print(f"  [warn] Docx not found: {doc_path} — skipping")
            continue

        print(f"\n  Processing {docx_name} → module='{module}'")
        try:
            doc = Document(str(doc_path))
        except Exception as exc:
            print(f"  [warn] Cannot open {docx_name}: {exc}")
            continue

        current_heading = "Overview"
        current_parts: list[str] = []
        chunks: list[dict[str, Any]] = []

        def flush(heading: str, parts: list[str]) -> None:
            text = "\n".join(parts).strip()
            if len(text) < MIN_CHUNK:
                return
            for sub in _chunk_by_size(text):
                chunks.append({
                    "text": f"[{module.upper()} Documentation / {heading}]\n{sub}",
                    "metadata": {
                        "module": module, "filename": docx_name,
                        "file_type": "documentation", "chunk_type": "doc_section",
                        "section": heading, "source": COLLECTION,
                    }
                })

        for para in doc.paragraphs:
            style = (para.style.name if para.style else "").lower()
            text = para.text.strip()
            if not text:
                continue
            if style.startswith("heading"):
                flush(current_heading, current_parts)
                current_heading = text
                current_parts = []
            else:
                current_parts.append(text)
        flush(current_heading, current_parts)

        # Tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                table_text = "\n".join(rows)
                if len(table_text) >= MIN_CHUNK:
                    for sub in _chunk_by_size(table_text):
                        chunks.append({
                            "text": f"[{module.upper()} Documentation / Table Data]\n{sub}",
                            "metadata": {
                                "module": module, "filename": docx_name,
                                "file_type": "documentation", "chunk_type": "doc_table",
                                "section": "Table Data", "source": COLLECTION,
                            }
                        })

        print(f"  {docx_name} → {len(chunks)} chunks")
        all_chunks.extend(chunks)

    return all_chunks


# ── Upsert to Qdrant ──────────────────────────────────────────────────────────
def upsert_chunks(client: QdrantClient, chunks: list[dict[str, Any]]) -> None:
    total = len(chunks)
    print(f"\n  Embedding and upserting {total} chunks to Qdrant...")

    BATCH_SIZE = 50
    points_batch: list[PointStruct] = []
    success = 0
    failed = 0

    for i, chunk in enumerate(chunks):
        text = chunk["text"]
        try:
            vector = openai_embed(text)
        except Exception as exc:
            print(f"  [warn] Embed failed for chunk {i}: {exc}")
            failed += 1
            time.sleep(1)
            continue

        points_batch.append(
            PointStruct(
                id=str(uuid.uuid4()),
                vector=vector,
                payload={**chunk["metadata"], "text": text},
            )
        )
        success += 1

        # Upsert in batches
        if len(points_batch) >= BATCH_SIZE:
            client.upsert(collection_name=COLLECTION, points=points_batch)
            points_batch = []
            print(f"  [{i+1}/{total}] Upserted batch ... ({success} ok, {failed} failed)")
            time.sleep(0.3)   # light rate-limit pause

    # Upsert remaining
    if points_batch:
        client.upsert(collection_name=COLLECTION, points=points_batch)
        print(f"  [{total}/{total}] Final batch upserted.")

    print(f"\n  ✅ Done: {success} embedded, {failed} failed out of {total} total chunks")


# ── Main ──────────────────────────────────────────────────────────────────────
def main() -> None:
    print("=" * 60)
    print("  QAD Custom Docs Embedder")
    print("  Collection:", COLLECTION)
    print("  Qdrant:    ", QDRANT_URL)
    print("  Embed model:", EMBED_MODEL)
    print("=" * 60)

    # Connect to Qdrant
    print("\n[1] Connecting to Qdrant...")
    client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=60)
    print("    Connected.")

    # Setup collection
    print("\n[2] Setting up collection...")
    setup_collection(client)

    # Collect chunks from source code ZIPs
    print("\n[3] Extracting source code from ZIPs...")
    code_chunks = collect_chunks_from_zips()
    print(f"    Total code chunks: {len(code_chunks)}")

    # Collect chunks from Word docs
    print("\n[4] Extracting content from Word documentation files...")
    doc_chunks = collect_chunks_from_docx()
    print(f"    Total doc chunks: {len(doc_chunks)}")

    all_chunks = code_chunks + doc_chunks
    print(f"\n    Grand total: {len(all_chunks)} chunks to embed")

    # Summary by module
    for mod in ["mrn", "doa", "rtdc"]:
        count = sum(1 for c in all_chunks if c["metadata"]["module"] == mod)
        print(f"      {mod.upper()}: {count} chunks")

    if not all_chunks:
        print("\n[ERROR] No chunks to embed. Check ZIP paths and file contents.")
        sys.exit(1)

    # Embed and upsert
    print("\n[5] Embedding and upserting to Qdrant...")
    upsert_chunks(client, all_chunks)

    # Final verification
    print("\n[6] Verifying collection...")
    info = client.get_collection(COLLECTION)
    print(f"    Points in collection: {info.points_count}")
    print("\n✅ All done! Apex Assistant can now answer questions about MRN, DOA, and RTDC.")
    print("   Make sure to select 'Custom Modules' domain in Apex Assistant.")


if __name__ == "__main__":
    main()
