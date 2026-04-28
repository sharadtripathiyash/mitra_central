"""
build_mrn_migration_blueprint.py
================================
One-shot renderer that produces `app/static/downloads/MRN_Migration_Blueprint.docx`.

This is a **pure renderer** — no LLM calls at runtime. All content below was
hand-curated from:
  - The existing MRN_System_Documentation.docx (custom MRN code analysis).
  - Web research on QAD Adaptive ERP native capabilities (Apr 2026).
  - Section 7 of the existing MRN doc that already lists native QAD overlap.

The output Word document is the second downloadable shown in the demo
Documentation tab (below the main system documentation), used only for the
MRN demo path. Re-run this script if you change the curated content below
or restyle the document.

Usage:
    python scripts/build_mrn_migration_blueprint.py
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

OUTPUT_PATH = Path("app/static/downloads/MRN_Migration_Blueprint.docx")

# ── Theme ─────────────────────────────────────────────────────────────────────
COLOR_INK         = RGBColor(0x0F, 0x1B, 0x3A)   # very dark blue (titles)
COLOR_TEAL        = RGBColor(0x00, 0x9B, 0x8C)   # teal accents
COLOR_BODY        = RGBColor(0x22, 0x2E, 0x4A)   # body text
COLOR_MUTED       = RGBColor(0x6B, 0x77, 0x90)   # captions
COLOR_TABLE_HEAD  = RGBColor(0x0F, 0x1B, 0x3A)
COLOR_TABLE_HEAD_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_HIGH        = RGBColor(0xC0, 0x39, 0x2B)
COLOR_MEDIUM      = RGBColor(0xC2, 0x8A, 0x00)
COLOR_LOW         = RGBColor(0x2F, 0x86, 0x4E)


# ── Low-level helpers ─────────────────────────────────────────────────────────
def _shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color="BFC8DC") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _run(paragraph, text: str, *, bold=False, size=10, color=COLOR_BODY,
         italic=False, font="Calibri") -> None:
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic


def _heading(doc, text: str, level: int = 1, color=COLOR_INK,
             space_before: int = 18, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    _run(p, text, bold=True, size=sizes.get(level, 11), color=color, font="Calibri")


def _para(doc, text: str, *, bold=False, italic=False, size=10.5, color=COLOR_BODY,
          space_after: int = 6, justify: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, text, bold=bold, italic=italic, size=size, color=color)


def _bullet(doc, text: str, *, level: int = 0, size=10) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(size)
        r.font.color.rgb = COLOR_BODY


def _kv_line(doc, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, f"{label}: ", bold=True, size=10, color=COLOR_INK)
    _run(p, value, size=10, color=COLOR_BODY)


def _table(doc, headers: list[str], rows: list[list[str]],
           col_widths: list[float] | None = None,
           row_shading_alt: str = "F4F7FB") -> None:
    """Build a styled table with header row + zebra-striped body."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        _shade(cell, "0F1B3A")
        _set_cell_borders(cell, color="0F1B3A")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _run(p, h, bold=True, size=10, color=COLOR_TABLE_HEAD_TEXT)

    # Body rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            if ri % 2 == 1 and row_shading_alt:
                _shade(cell, row_shading_alt)
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Detect High / Medium / Low badges in the last column for colourisation
            color = COLOR_BODY
            stripped = val.strip().lower()
            if stripped in ("high", "high effort"):
                color = COLOR_HIGH
            elif stripped in ("medium", "med", "medium effort"):
                color = COLOR_MEDIUM
            elif stripped in ("low", "low effort"):
                color = COLOR_LOW
            _run(p, val, size=9.5, color=color)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)


def _code_block(doc, lines: Iterable[str]) -> None:
    """Render a SQL/DDL block in a shaded monospace box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.1)

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F3F8")
    pPr.append(shd)

    # Borders on the paragraph
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "BFC8DC")
        pBdr.append(b)
    pPr.append(pBdr)

    text = "\n".join(lines)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = COLOR_INK
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def _hrule(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "BFC8DC")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Content ───────────────────────────────────────────────────────────────────
def _cover(doc) -> None:
    # Tag line above title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(8)
    _run(p, "QAD ADAPTIVE MIGRATION BLUEPRINT", bold=True, size=11, color=COLOR_TEAL)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _run(p, "Migration Blueprint:", bold=True, size=26, color=COLOR_INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    _run(p, "Bridging the QAD Adaptive Gap", bold=True, size=26, color=COLOR_INK)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    _run(
        p,
        "MRN — From Custom Progress 4GL to Native QAD Adaptive Implementation",
        italic=True, size=12, color=COLOR_MUTED,
    )

    # Companion-doc note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _run(
        p,
        "Companion to MRN_System_Documentation.docx",
        italic=True, size=10, color=COLOR_MUTED,
    )

    # Metadata block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Module: ", bold=True, size=10, color=COLOR_INK)
    _run(p, "MRN — Material Requisition Note System", size=10, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Source platform: ", bold=True, size=10, color=COLOR_INK)
    _run(p, "QAD ERP | Progress 4GL custom development", size=10, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Target platform: ", bold=True, size=10, color=COLOR_INK)
    _run(p, "QAD Adaptive ERP with Adaptive UX (2022.1+)", size=10, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Audience: ", bold=True, size=10, color=COLOR_INK)
    _run(p, "Business stakeholders + QAD Adaptive developers", size=10, color=COLOR_BODY)

    doc.add_page_break()


def _section_1_executive_summary(doc) -> None:
    _heading(doc, "1. Executive Summary", level=1, space_before=0)

    _para(
        doc,
        "The Material Requisition Note (MRN) system is a custom Progress 4GL application built on QAD ERP. "
        "Internal capability analysis against QAD Adaptive ERP (2022.1+) shows that approximately 65 % of MRN's "
        "business capabilities are already covered natively by standard QAD Adaptive modules — Inventory Control "
        "miscellaneous transactions, the .NET UX inventory transfer engine, lot/serial tracking, and product-line "
        "based GL determination. The remaining ~35 % represents purpose-driven business logic that is unique to the "
        "internal-store-to-department workflow this system was originally built to support, and is not available "
        "out of the box in QAD Adaptive.",
    )

    _para(
        doc,
        "Closing this gap on QAD Adaptive ERP requires a structured customisation effort spanning three layers: "
        "(1) eight new custom database tables for purpose codes, requisition types and the audit trail; "
        "(2) five OOABL Business Components implementing the approval engine, GL resolver, validity-date check, "
        "transfer validator and execution router; (3) four Adaptive UX TypeScript screens that recreate the "
        "data-entry, approval, execution and inquiry experience using the QAD WebUI extension framework. The work "
        "integrates into standard QAD by hooking pre/post triggers on the IC Misc transaction APIs (icintr.p, "
        "icxfer.p) and the existing GL/lot/serial subsystems — meaning the custom work adds workflow on top of "
        "standard transactions, rather than reimplementing them.",
    )

    _para(
        doc,
        "Total effort is estimated at 16–22 person-weeks of QAD-experienced development across one back-end OOABL "
        "developer, one Adaptive UX (TypeScript) developer, one functional analyst and a part-time test lead. The "
        "recommended approach is a phased delivery: Phase 1 builds the masters and approval engine (~6 weeks), "
        "Phase 2 builds the transaction execution and GL resolver (~7 weeks), Phase 3 builds the inquiries, audit "
        "history and cutover (~5 weeks). All Phase-1 deliverables are usable on their own — the team can run a "
        "pilot site after Phase 1 even before transactions are wired in.",
    )

    _para(
        doc,
        "This blueprint is intended to be read alongside MRN_System_Documentation.docx. The companion document "
        "explains what the existing custom code does; this document explains what to build inside QAD Adaptive ERP "
        "to replace it.",
        italic=True, color=COLOR_MUTED,
    )


def _section_2_gap_analysis(doc) -> None:
    _heading(doc, "2. Capability-by-Capability Gap Analysis", level=1)

    _para(
        doc,
        "Each row below maps one MRN business capability to its closest standard QAD Adaptive equivalent and "
        "describes the residual gap that must be bridged through customisation. Coverage values are based on "
        "QAD Adaptive ERP 2022.1 with Adaptive UX. Items marked Full are entirely satisfied by standard QAD; "
        "Partial items are satisfied in part but require an extension; None items have no native equivalent and "
        "must be built from scratch.",
    )

    _table(
        doc,
        headers=["Custom Capability", "QAD Adaptive Coverage", "Native Module / Feature", "Gap to Bridge"],
        rows=[
            [
                "Raise internal material requisitions against configurable types and sites",
                "Partial",
                "IC Misc Transactions (icintr.p) + Site/Location Master",
                "No internal-requisition document exists. Need a custom requisition header/detail entity that wraps standard IC Misc as the execution layer.",
            ],
            [
                "Multi-level authorisation: Create → Approve → Execute with user-group restrictions",
                "Partial",
                "QAD Workflow Engine + code_mstr user groups",
                "Native workflow does role notification but cannot enforce a Create→Approve→Execute state machine on a custom document. Need a custom approval-state engine plus group gating per state.",
            ],
            [
                "Purpose-based GL account determination (product-line OR direct account)",
                "Partial",
                "Product Line Master + Account Master + sub-account/cost-centre fields",
                "QAD resolves accounts from product line or transaction reason. There is no purpose-code abstraction that can override either. Need a Purpose Master plus a resolver that runs before posting.",
            ],
            [
                "Lot/serial control during execution (full / same-lot / new-lot modes)",
                "Partial",
                "IC Lot/Serial control flag at item level + standard lot prompts",
                "QAD enforces lot tracking on/off per item but has no per-requisition-type lot mode (same-lot vs new-lot). Need a per-requisition-type LTC override layer.",
            ],
            [
                "Inter-site / inter-location inventory transfer with conflict checks",
                "Full",
                "IC Transfer (icxfer.p) + maintainInventoryTransfer API",
                "Native covers the transfer mechanics. Custom validation rules (e.g. blocking same-site/same-loc transfers, validating the requisition's purpose against the destination) live in the wrapper.",
            ],
            [
                "Validity-date enforcement — expired requisition lines cannot be executed",
                "None",
                "Not available natively",
                "Must be implemented as a pre-execute validation hook on the custom requisition detail entity.",
            ],
            [
                "Full audit trail of every state change (Create / Submit / Approve / Reject / Execute)",
                "Partial",
                "Standard tr_hist (transaction) + audit log on user fields",
                "tr_hist captures inventory movements but not approval-state transitions. Need a custom history table mirroring xxmrh_hist.",
            ],
            [
                "Inventory availability display before transaction entry",
                "Full",
                "Adaptive UX Inventory Inquiry widget + browse APIs",
                "Reuse the standard Inventory Inquiry component — embed it in the custom requisition entry screen via the Adaptive UX composition framework.",
            ],
        ],
        col_widths=[1.5, 1.0, 1.7, 2.4],
    )

    _para(
        doc,
        "Out of eight major capabilities, two are fully covered by native QAD, five are partially covered (need a "
        "custom layer that integrates with the standard module), and one (validity-date enforcement) has no native "
        "equivalent and must be built from scratch. This roughly corresponds to the 65 % replaceability score shown "
        "on the analysis dashboard — 65 % of the work is leveraging native QAD, 35 % is the custom bridge described "
        "in the rest of this document.",
        italic=True,
        color=COLOR_MUTED,
    )


def _section_3_custom_tables(doc) -> None:
    _heading(doc, "3. Custom Tables to Build", level=1)
    _para(
        doc,
        "The following eight tables must be created inside the QAD Adaptive ERP database using Application Builder "
        "or via standard Progress data-dictionary scripts (.df). All tables follow the QAD naming convention "
        "(xx-prefix for custom). Each table is multi-domain by including dom_part_id (or equivalent) in the "
        "primary index, matching how standard QAD partitions data by domain. Field types use Progress data-types "
        "(character, integer, decimal, date, logical, datetime-tz). Foreign keys to standard QAD tables are not "
        "enforced at the database level (Progress convention) but must be enforced in the OOABL Business Components.",
    )

    tables = [
        (
            "3.1  xxmrc_ctrl — MRN Control File",
            "Single-row control table holding system-wide MRN settings.",
            [
                ("xxmrc_domain",       "char(8)",      "Domain code (PK)"),
                ("xxmrc_next_nbr",     "integer",      "Next available requisition number"),
                ("xxmrc_default_site", "char(8)",      "Default site for new requisitions"),
                ("xxmrc_dflt_validity_days", "integer", "Default validity period (days from creation)"),
                ("xxmrc_email_on_submit", "logical",   "Whether to email approvers on submit"),
                ("xxmrc_user",         "char(20)",     "Last modified by"),
                ("xxmrc_mod_date",     "date",         "Last modified date"),
            ],
            "Unique index: xxmrc_domain",
        ),
        (
            "3.2  xxmt_mstr — Requisition Type Master",
            "Defines requisition types (e.g. ISSUE, RECEIPT, TRANSFER) and their default behaviour.",
            [
                ("xxmt_domain",   "char(8)",  "Domain code (PK)"),
                ("xxmt_code",     "char(8)",  "Requisition type code (PK)"),
                ("xxmt_desc",     "char(60)", "Description"),
                ("xxmt_trx_type", "char(10)", "Underlying QAD transaction type (ISS-UNP / RCT-UNP / ISS-TR / RCT-TR)"),
                ("xxmt_trx_ltc",  "char(10)", "Lot/serial control: FULL / SAME / NEW / NONE"),
                ("xxmt_create_grp", "char(8)", "code_mstr group authorised to create"),
                ("xxmt_approve_grp", "char(8)", "code_mstr group authorised to approve"),
                ("xxmt_execute_grp", "char(8)", "code_mstr group authorised to execute"),
                ("xxmt_active",   "logical",  "Active flag"),
            ],
            "Unique index: xxmt_domain + xxmt_code",
        ),
        (
            "3.3  xxmtd_det — Requisition Type → Allowed Purposes",
            "Many-to-many bridge constraining which purpose codes are valid for each requisition type.",
            [
                ("xxmtd_domain",  "char(8)", "Domain code (PK)"),
                ("xxmtd_type",    "char(8)", "Requisition type code (PK, FK → xxmt_mstr)"),
                ("xxmtd_purpose", "char(8)", "Purpose code (PK, FK → xxmp_mstr)"),
                ("xxmtd_default", "logical", "Whether this is the default purpose for the type"),
            ],
            "Unique index: xxmtd_domain + xxmtd_type + xxmtd_purpose",
        ),
        (
            "3.4  xxmp_mstr — Purpose Code Master",
            "Defines purposes that drive GL-account determination at execution time.",
            [
                ("xxmp_domain",  "char(8)",     "Domain code (PK)"),
                ("xxmp_code",    "char(8)",     "Purpose code (PK)"),
                ("xxmp_desc",    "char(60)",    "Description"),
                ("xxmp_pa",      "char(1)",     "Account source: P = product line, A = direct account"),
                ("xxmp_account", "char(20)",    "Direct account (if xxmp_pa = A)"),
                ("xxmp_sub",     "char(20)",    "Sub-account (if xxmp_pa = A)"),
                ("xxmp_cc",      "char(20)",    "Cost centre (if xxmp_pa = A)"),
                ("xxmp_active",  "logical",     "Active flag"),
            ],
            "Unique index: xxmp_domain + xxmp_code",
        ),
        (
            "3.5  xxmpd_det — Purpose Valid Part Ranges",
            "Restricts which item parts a purpose code applies to (range-based).",
            [
                ("xxmpd_domain",  "char(8)",  "Domain code (PK)"),
                ("xxmpd_purpose", "char(8)",  "Purpose code (PK, FK → xxmp_mstr)"),
                ("xxmpd_seq",     "integer",  "Range sequence (PK)"),
                ("xxmpd_part_lo", "char(18)", "Part number range start"),
                ("xxmpd_part_hi", "char(18)", "Part number range end"),
                ("xxmpd_pl",      "char(8)",  "Optional product line override"),
            ],
            "Unique index: xxmpd_domain + xxmpd_purpose + xxmpd_seq",
        ),
        (
            "3.6  xxmr_mstr — Requisition Header",
            "Header for one requisition document.",
            [
                ("xxmr_domain",  "char(8)",   "Domain code (PK)"),
                ("xxmr_nbr",     "integer",   "Requisition number (PK)"),
                ("xxmr_type",    "char(8)",   "Requisition type (FK → xxmt_mstr)"),
                ("xxmr_site",    "char(8)",   "Site code (FK → si_mstr)"),
                ("xxmr_purpose", "char(8)",   "Purpose code (FK → xxmp_mstr)"),
                ("xxmr_status",  "char(10)",  "DRAFT / SUBMITTED / APPROVED / REJECTED / EXECUTED / CANCELLED"),
                ("xxmr_validity_dt", "date",  "Validity expiry date"),
                ("xxmr_create_user", "char(20)", "Creator user-id"),
                ("xxmr_create_dt", "datetime-tz", "Created timestamp"),
                ("xxmr_approve_user", "char(20)", "Approver user-id (when status = APPROVED)"),
                ("xxmr_approve_dt", "datetime-tz", "Approval timestamp"),
                ("xxmr_remarks", "char(255)", "Free-text remarks"),
            ],
            "Unique index: xxmr_domain + xxmr_nbr; secondary: xxmr_status + xxmr_create_dt",
        ),
        (
            "3.7  xxmrd_det — Requisition Detail Lines",
            "Line items within a requisition.",
            [
                ("xxmrd_domain", "char(8)",  "Domain code (PK)"),
                ("xxmrd_nbr",    "integer",  "Requisition number (PK, FK → xxmr_mstr)"),
                ("xxmrd_line",   "integer",  "Line sequence (PK)"),
                ("xxmrd_part",   "char(18)", "Item number (FK → pt_mstr)"),
                ("xxmrd_qty",    "decimal",  "Quantity requested"),
                ("xxmrd_um",     "char(2)",  "Unit of measure"),
                ("xxmrd_loc_from", "char(8)", "Source location"),
                ("xxmrd_loc_to",   "char(8)", "Destination location (transfers only)"),
                ("xxmrd_lot",    "char(18)", "Lot/serial reference (when LTC=SAME)"),
                ("xxmrd_validity_dt", "date", "Per-line validity expiry"),
                ("xxmrd_executed", "logical", "TRUE once posted to inventory"),
                ("xxmrd_tr_nbr", "integer",  "Reference to standard tr_hist row created on execution"),
            ],
            "Unique index: xxmrd_domain + xxmrd_nbr + xxmrd_line",
        ),
        (
            "3.8  xxmrh_hist — Requisition State-Change Audit",
            "Append-only history of every approval/execution state change. Mirrors how tr_hist works for inventory.",
            [
                ("xxmrh_domain",  "char(8)",   "Domain code"),
                ("xxmrh_nbr",     "integer",   "Requisition number"),
                ("xxmrh_seq",     "integer",   "Event sequence (PK)"),
                ("xxmrh_action",  "char(15)",  "CREATE / SUBMIT / APPROVE / REJECT / EXECUTE / CANCEL"),
                ("xxmrh_old_status", "char(10)", "Status before action"),
                ("xxmrh_new_status", "char(10)", "Status after action"),
                ("xxmrh_user",    "char(20)",  "User performing the action"),
                ("xxmrh_dt",      "datetime-tz", "Timestamp"),
                ("xxmrh_comment", "char(255)", "Optional comment from approver/rejecter"),
            ],
            "Unique index: xxmrh_domain + xxmrh_nbr + xxmrh_seq",
        ),
    ]

    for title, desc, fields, idx in tables:
        _heading(doc, title, level=2)
        _para(doc, desc, italic=True, color=COLOR_MUTED)
        _table(
            doc,
            headers=["Field", "Type", "Description"],
            rows=[[f, t, d] for f, t, d in fields],
            col_widths=[1.6, 1.4, 3.6],
        )
        _para(doc, idx, size=9.5, color=COLOR_MUTED, justify=False, space_after=10)


def _section_4_business_components(doc) -> None:
    _heading(doc, "4. Business Components to Build", level=1)
    _para(
        doc,
        "Five OOABL Business Components form the back-end of the bridge layer. Each is a server-side Progress "
        "OpenEdge ABL component exposed to the Adaptive UX front-end through the standard QAD WebUI REST API. "
        "Each component should be built using the QAD Application Builder so that it inherits standard plumbing "
        "(security, logging, multi-domain, transactional integrity). Naming follows the bcXxMrn* convention used "
        "by QAD's own components.",
    )

    _table(
        doc,
        headers=["Component", "Type", "Purpose / Hook Points", "Effort"],
        rows=[
            [
                "bcXxMrnRequisition",
                "Document Business Component",
                "Owns CRUD + state transitions for xxmr_mstr / xxmrd_det. Exposes operations: createDraft, addLine, submit, approve, reject, execute, cancel. Each transition writes an xxmrh_hist row. Calls bcXxMrnApproval.canTransition() before any state change. Triggers email notifications on submit/approve/reject (uses standard QAD email API).",
                "High",
            ],
            [
                "bcXxMrnApproval",
                "Workflow Component",
                "Centralised approval-state engine. Encapsulates the Create→Approve→Execute state machine: enforces user-group gating per state by reading xxmt_mstr.xxmt_create_grp/_approve_grp/_execute_grp and matching the current user's code_mstr groups. Single point of policy — replaces the ad-hoc approval logic scattered across MRN's legacy programs.",
                "High",
            ],
            [
                "bcXxMrnGlResolver",
                "Service Component",
                "Resolves the GL account for a posting line. Reads xxmp_mstr.xxmp_pa: when 'A', returns xxmp_account/_sub/_cc directly; when 'P', falls through to the standard product-line account on pl_mstr for the part's product line. Plugs into the IC Misc post-write hook (icintr.p override) so all postings — custom or standard — go through the same resolver.",
                "Medium",
            ],
            [
                "bcXxMrnValidityCheck",
                "Validation Component",
                "Pre-execute hook. Reads xxmrd_det.xxmrd_validity_dt and rejects execution if today > validity. Also compares the requisition header's xxmr_validity_dt and uses the more restrictive of the two. Returns a structured error the Adaptive UX layer can surface as a field-level message.",
                "Low",
            ],
            [
                "bcXxMrnTransferValidator",
                "Validation Component",
                "Pre-execute hook for transfer-type requisitions. Validates: (a) source ≠ destination location, (b) destination location is active and accepts the part's product line, (c) the transfer does not violate any standard QAD inventory location-restriction rules. Composes onto bcXxMrnRequisition.execute() before the call to maintainInventoryTransfer.",
                "Low",
            ],
        ],
        col_widths=[1.5, 1.2, 3.5, 0.8],
    )

    _heading(doc, "Integration with Standard QAD Transactions", level=2)
    _para(
        doc,
        "All four execution flows wrap standard QAD transaction APIs rather than reimplementing them — this is "
        "the single most important architectural rule for keeping the customisation upgrade-safe:",
    )
    _bullet(doc, "ISS-UNP / RCT-UNP (issues and unplanned receipts) → call icintr.p")
    _bullet(doc, "ISS-TR / RCT-TR (inventory transfers) → call maintainInventoryTransfer (Adaptive UX) or icxfer.p (legacy)")
    _bullet(doc, "GL posting → reuse icintra.p which already understands product-line accounts; bcXxMrnGlResolver overrides the account before the call")
    _bullet(doc, "Lot/serial prompts → reuse standard lot prompt component; bcXxMrnRequisition passes the per-type LTC mode (FULL / SAME / NEW / NONE) as input")


def _section_5_ux_screens(doc) -> None:
    _heading(doc, "5. UI / Screen Components", level=1)
    _para(
        doc,
        "The user-facing layer consists of four Adaptive UX (TypeScript) screens built using the QAD WebUI "
        "extension framework. Each screen is a TypeScript class that consumes the corresponding back-end Business "
        "Component via the auto-generated REST API. Standard QAD components (browse, lookup, inventory-availability "
        "widget) are composed in rather than reimplemented.",
    )

    screens = [
        (
            "5.1  Requisition Maintenance (xxMrnRequisitionMaint)",
            "Primary data-entry screen. Roughly equivalent to the legacy xxmrmt.p + xxmrmt01.p frames combined.",
            [
                "Header section: requisition type (lookup → xxmt_mstr), site (lookup → si_mstr), purpose (lookup → xxmp_mstr filtered by xxmtd_det), validity date (date-picker, defaults from xxmrc_ctrl)",
                "Detail grid: editable line table with part lookup (pt_mstr), qty/UM, source location, destination location (shown only for transfer types), per-line validity",
                "Embedded Inventory Availability widget: standard QAD component, refreshes when part + site change",
                "Action bar: Save Draft, Submit for Approval, Cancel, Print",
                "All field-level validations come from bcXxMrnRequisition; UI only renders the structured errors",
            ],
        ),
        (
            "5.2  Approval Inbox (xxMrnApprovalInbox)",
            "Role-filtered queue of requisitions awaiting the current user's approval.",
            [
                "Filtered list: WHERE xxmr_status = SUBMITTED AND user is in xxmt_mstr.xxmt_approve_grp",
                "Click-through to a read-only Requisition view with two action buttons: Approve, Reject (with mandatory comment)",
                "Each action calls bcXxMrnApproval.approve() or .reject() — which in turn writes xxmrh_hist and triggers the email notification",
                "Bulk-approve checkbox UI for low-risk requisition types (configurable on xxmt_mstr)",
            ],
        ),
        (
            "5.3  Execution Screen (xxMrnExecutionScreen)",
            "Where approved requisitions are posted to inventory. Different forms for issue/receipt vs transfer.",
            [
                "Lists APPROVED requisitions where the user is in xxmt_mstr.xxmt_execute_grp",
                "Per-line lot/serial prompt — driven by xxmt_mstr.xxmt_trx_ltc (FULL prompts for every unit, SAME prompts once per line, NEW auto-generates a lot, NONE skips)",
                "On Post: bcXxMrnRequisition.execute() runs the validity + transfer validators, then calls icintr.p / icxfer.p, then writes xxmrh_hist, then updates xxmrd_det.xxmrd_executed and xxmrd_tr_nbr with the standard tr_hist reference",
                "Standard QAD posting errors are surfaced inline and the requisition stays on screen for retry",
            ],
        ),
        (
            "5.4  Inquiry / Reporting (xxMrnInquiry)",
            "Read-only browser for completed and in-flight requisitions plus history drill-down.",
            [
                "Filters: status, type, site, purpose, date range, requester, approver",
                "Drill-down on a row opens the full requisition with a History panel rendering the xxmrh_hist timeline",
                "Export to Excel via the standard QAD WebUI export hook",
                "Companion print report (xxmrPrintReport) replaces the legacy XXMRPMQ1.P inquiry report",
            ],
        ),
    ]
    for title, desc, items in screens:
        _heading(doc, title, level=2)
        _para(doc, desc, italic=True, color=COLOR_MUTED, space_after=4)
        for it in items:
            _bullet(doc, it)


def _section_6_workflow_integration(doc) -> None:
    _heading(doc, "6. Workflow & Integration Extensions", level=1)

    _heading(doc, "6.1  Approval State Machine", level=2)
    _para(
        doc,
        "QAD Adaptive ERP's native Workflow Engine handles role-based notifications well but is weak at enforcing "
        "a custom document state machine. The recommended pattern is to keep all state-transition rules inside "
        "bcXxMrnApproval (Section 4) and use QAD Workflow only for the email/notification fan-out. The state "
        "machine is:",
    )
    _code_block(doc, [
        "[DRAFT] --submit--> [SUBMITTED]",
        "[SUBMITTED] --approve--> [APPROVED]",
        "[SUBMITTED] --reject--> [REJECTED] (terminal)",
        "[APPROVED] --execute--> [EXECUTED] (terminal)",
        "[DRAFT|SUBMITTED|APPROVED] --cancel--> [CANCELLED] (terminal)",
        "",
        "Guard on every transition:",
        "  user must belong to the code_mstr group named in xxmt_mstr",
        "  for the from-state (e.g. xxmt_approve_grp gates SUBMITTED -> APPROVED).",
    ])

    _heading(doc, "6.2  Email Notification Hooks", level=2)
    _para(
        doc,
        "The standard QAD email API (qadEmailSend.p) is invoked by bcXxMrnRequisition on three transitions: "
        "(a) submit — emails every user in xxmt_approve_grp with a deep-link to the Approval Inbox; "
        "(b) approve — emails the original requester confirming approval; "
        "(c) reject — emails the requester with the rejection comment. Email templates live in the standard "
        "xxXxMrn*.html templates folder shipped alongside the customisation package.",
    )

    _heading(doc, "6.3  Pre/Post Triggers on Standard QAD Transactions", level=2)
    _para(
        doc,
        "bcXxMrnGlResolver is registered as a pre-write trigger on the IC Misc transaction APIs. The QAD "
        "extension framework allows pre/post hooks via the standard customisation registry:",
    )
    _code_block(doc, [
        "/* Register in xxmrCustomTrig.i */",
        "RUN registerTrigger IN h_extReg",
        "    ('icintr.p', 'PRE-WRITE', 'bcXxMrnGlResolver.resolve').",
        "RUN registerTrigger IN h_extReg",
        "    ('maintainInventoryTransfer', 'PRE-EXECUTE', 'bcXxMrnTransferValidator.validate').",
    ])

    _heading(doc, "6.4  Future Java Migration Note", level=2)
    _para(
        doc,
        "QAD has signalled that the Adaptive ERP back-end will migrate from Progress OpenEdge to Java in a future "
        "release. To keep this customisation upgrade-safe: keep all business logic inside the named OOABL Business "
        "Components (which will be rewritten in Java by QAD using the same component contract), and avoid embedding "
        "logic inside trigger procedures or include files. The TypeScript front-end is unaffected by the Java "
        "migration.",
        italic=True, color=COLOR_MUTED,
    )


def _section_7_roadmap(doc) -> None:
    _heading(doc, "7. Phased Implementation Roadmap", level=1)
    _para(
        doc,
        "Three phases. Each phase produces a usable, testable artefact — a pilot site can run on Phase 1 alone "
        "for proof-of-concept before transactional posting is enabled in Phase 2.",
    )

    _heading(doc, "Phase 1 — Masters & Approval Engine (~6 weeks)", level=2)
    _table(
        doc,
        headers=["Deliverable", "Component", "Effort"],
        rows=[
            ["8 custom tables created via Application Builder .df scripts",
             "DB schema (Section 3)", "1 week"],
            ["xxmrc_ctrl, xxmt_mstr, xxmtd_det, xxmp_mstr, xxmpd_det maintenance screens",
             "Adaptive UX screens", "2 weeks"],
            ["bcXxMrnRequisition + bcXxMrnApproval (state machine, no execute yet)",
             "Business Components", "2 weeks"],
            ["Requisition Maintenance + Approval Inbox screens (Sections 5.1, 5.2)",
             "Adaptive UX screens", "1 week"],
        ],
        col_widths=[3.4, 2.4, 1.0],
    )
    _para(doc, "Phase 1 exit: users can create, submit and approve requisitions end-to-end. No inventory impact yet — approved requisitions stay in APPROVED state.",
          italic=True, color=COLOR_MUTED)

    _heading(doc, "Phase 2 — Execution & GL Resolver (~7 weeks)", level=2)
    _table(
        doc,
        headers=["Deliverable", "Component", "Effort"],
        rows=[
            ["bcXxMrnGlResolver wired as IC Misc pre-write trigger",
             "Business Component + trigger registration", "1.5 weeks"],
            ["bcXxMrnValidityCheck + bcXxMrnTransferValidator",
             "Business Components", "1 week"],
            ["Execution screen with lot/serial prompts (Section 5.3)",
             "Adaptive UX screen", "2 weeks"],
            ["Email notification hooks (Section 6.2)",
             "Email templates + integration", "0.5 week"],
            ["Integration test pack: 4 transaction types × lot modes × purpose codes",
             "Test scripts + harness", "2 weeks"],
        ],
        col_widths=[3.4, 2.4, 1.0],
    )
    _para(doc, "Phase 2 exit: approved requisitions can be posted to inventory through standard QAD APIs. tr_hist captures every movement; xxmrh_hist captures every approval.",
          italic=True, color=COLOR_MUTED)

    _heading(doc, "Phase 3 — Inquiry, Audit & Cutover (~5 weeks)", level=2)
    _table(
        doc,
        headers=["Deliverable", "Component", "Effort"],
        rows=[
            ["Inquiry screen with history timeline (Section 5.4)",
             "Adaptive UX screen", "1.5 weeks"],
            ["Print report — xxmrPrintReport (replaces XXMRPMQ1.P)",
             "Report definition + template", "0.5 week"],
            ["Data migration scripts: legacy xxmr_mstr / xxmrd_det → new tables",
             "Migration scripts + verification", "1.5 weeks"],
            ["UAT support + parallel run (1 week production-shadow)",
             "Test + cutover", "1 week"],
            ["Go-live, post-go-live hypercare",
             "Cutover + support", "0.5 week"],
        ],
        col_widths=[3.4, 2.4, 1.0],
    )
    _para(doc, "Phase 3 exit: legacy custom MRN code is decommissioned; all flow goes through the new Adaptive UX layer; full historical data preserved.",
          italic=True, color=COLOR_MUTED)


def _section_8_total_estimate(doc) -> None:
    _heading(doc, "8. Total Build Estimate", level=1)

    _para(doc,
          "The complete bridge requires the following deliverables. Effort is in person-weeks at QAD-experienced "
          "developer rates; calendar time depends on team size and parallelisation.",
          space_after=10)

    _table(
        doc,
        headers=["Category", "Items", "Total"],
        rows=[
            ["Custom database tables (Section 3)",
             "8 tables (xxmrc_ctrl, xxmt_mstr, xxmtd_det, xxmp_mstr, xxmpd_det, xxmr_mstr, xxmrd_det, xxmrh_hist)",
             "1 week"],
            ["OOABL Business Components (Section 4)",
             "5 components: Requisition, Approval, GlResolver, ValidityCheck, TransferValidator",
             "5.5 weeks"],
            ["Adaptive UX TypeScript screens (Section 5)",
             "4 screens: Maintenance, Approval Inbox, Execution, Inquiry",
             "5 weeks"],
            ["Workflow / email integration (Section 6)",
             "1 state machine, 3 email templates, 2 trigger registrations",
             "1 week"],
            ["Migration scripts + test harness",
             "Data migration + 4×LTC×purpose integration tests",
             "3.5 weeks"],
            ["UAT, parallel run, cutover (Phase 3)",
             "1-week production shadow + go-live support",
             "1.5 weeks"],
        ],
        col_widths=[2.4, 3.4, 1.0],
    )

    _heading(doc, "Headline Estimate", level=2)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Total effort: ", bold=True, size=11, color=COLOR_INK)
    _run(p, "16–22 person-weeks (≈ 4–5 calendar months with a 2-developer team)",
         size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Team composition: ", bold=True, size=11, color=COLOR_INK)
    _run(p, "1 OOABL/Progress developer · 1 Adaptive UX (TypeScript) developer · "
            "1 functional analyst · 0.5 FTE test lead", size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Risk band: ", bold=True, size=11, color=COLOR_INK)
    _run(p, "Medium — the integration with standard IC Misc and IC Transfer is well-trodden; "
            "the residual risk is the Adaptive UX framework's evolving extension API and the "
            "future Progress→Java back-end migration.", size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Recommendation: ", bold=True, size=11, color=COLOR_INK)
    _run(p, "Proceed in three phases as described in Section 7. The Phase-1 deliverable is "
            "independently valuable and lets the business validate the data model and approval flow "
            "before committing to the full execution layer.",
         size=11, color=COLOR_BODY)

    _hrule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    _run(p, "End of Migration Blueprint", italic=True, size=10, color=COLOR_MUTED)


# ── Top-level orchestrator ────────────────────────────────────────────────────
def build() -> Path:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default font for the whole doc
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _cover(doc)
    _section_1_executive_summary(doc)
    _section_2_gap_analysis(doc)
    _section_3_custom_tables(doc)
    _section_4_business_components(doc)
    _section_5_ux_screens(doc)
    _section_6_workflow_integration(doc)
    _section_7_roadmap(doc)
    _section_8_total_estimate(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = build()
    print(f"Wrote {out}  ({out.stat().st_size:,} bytes)")
